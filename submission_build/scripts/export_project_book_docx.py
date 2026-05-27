from pathlib import Path
from docx import Document
from docx.shared import Pt

root = Path(__file__).resolve().parents[1]
md_path = root / "docs" / "PROJECT_BOOK_HE.md"
docx_path = root / "docs" / "PROJECT_BOOK_HE.docx"

lines = md_path.read_text(encoding="utf-8").splitlines()

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(12)

for raw in lines:
    line = raw.rstrip()
    if not line:
        doc.add_paragraph("")
        continue

    if line.startswith("### "):
        doc.add_heading(line[4:], level=3)
        continue
    if line.startswith("## "):
        doc.add_heading(line[3:], level=2)
        continue
    if line.startswith("# "):
        doc.add_heading(line[2:], level=1)
        continue

    if line.startswith("- "):
        p = doc.add_paragraph(line[2:], style="List Bullet")
        p.style.font.size = Pt(12)
        continue

    if line[:2].isdigit() and line[1:3] == ". ":
        doc.add_paragraph(line[3:], style="List Number")
        continue

    if line.startswith("---"):
        doc.add_paragraph("")
        continue

    doc.add_paragraph(line)

section = doc.sections[0]
section.header.paragraphs[0].text = "יוסף יצחק משלחיס | Trivia Game"
section.footer.paragraphs[0].text = "עמוד"

try:
    doc.save(docx_path)
    print(docx_path)
except PermissionError:
    alt = docx_path.with_name("PROJECT_BOOK_HE_FIXED.docx")
    doc.save(alt)
    print(alt)
 
